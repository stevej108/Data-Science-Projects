#------------------
# Agent 7 Helper Functions
#----------------------------

from docx import Document
from docx.shared import Pt

# ==========================================================
# Word Hyperlink Helper
# ==========================================================

from pathlib import Path
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from core.logger import get_logger

logger = get_logger(__name__)


def add_hyperlink(paragraph, text, target):
    """
    Adds a clickable hyperlink to a Word paragraph.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
    text : str
        Display text
    target : str
        Local path or URL
    """

    # Convert local file paths to file:// URI
    if not str(target).startswith(("http://", "https://", "file://")):
        target = Path(target).resolve().as_uri()

    part = paragraph.part

    r_id = part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text

    new_run.append(text_element)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def create_memo_doc(
    query: str,
    executive_summary: str,
    key_findings: list,
    recommendations: list,
    document_summaries: list
):

    doc = Document()

    # ==========================================================
    # Title
    # ==========================================================
    doc.add_heading("Query Summary", level=1)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph("")

    # ==========================================================
    # Table of Contents
    # ==========================================================
    doc.add_heading("Table of Contents", level=2)

    toc_items = [
        "Executive Summary",
        "Key Findings",
        "Supporting Insights",
        "Source Appendix"
    ]

    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph("")

    # ==========================================================
    # Executive Summary
    # ==========================================================
    doc.add_heading("Executive Summary", level=2)

    doc.add_paragraph(executive_summary or "No summary available.")

    doc.add_paragraph("")

    # ==========================================================
    # Key Findings
    # ==========================================================
    doc.add_heading("Key Findings", level=2)

    for i, finding in enumerate(key_findings or [], start=1):

        p = doc.add_paragraph()

        p.add_run(f"{i}. {finding.get('title', 'Untitled')}").bold = True

        p.add_run(f"\n{finding.get('insight', '')}")

        if finding.get("business_impact"):
            p.add_run("\nBusiness Impact: ").bold = True
            p.add_run(finding["business_impact"])

        if finding.get("supporting_documents"):
            p.add_run("\nSupporting Documents: ").bold = True
            p.add_run(", ".join(finding["supporting_documents"]))

        doc.add_paragraph("")

    # ==========================================================
    # Supporting Insights
    # ==========================================================
    doc.add_heading("Supporting Insights", level=2)

    for doc_item in document_summaries or []:

        doc.add_paragraph(
            f"• {doc_item.get('file_name','Unknown')} | "
            f"{doc_item.get('relative_path','')}"
        )

        doc.add_paragraph(doc_item.get("summary", ""))

        doc.add_paragraph("")

    # ==========================================================
    # References
    # ==========================================================
    from collections import defaultdict

    doc.add_heading("References", level=1)

    documents = defaultdict(list)

    for doc_item in document_summaries or []:
        for src in doc_item.get("sources", []):
            documents[doc_item["file_name"]].append(src)

    for file_name, chunks in sorted(documents.items()):

        if not chunks:
            continue

        first = chunks[0]

        p = doc.add_paragraph()

        p.add_run(file_name).bold = True
        p.add_run(f"\nType: {first.get('document_type','')}")
        p.add_run(f"\nFolder: {first.get('relative_path','')}")

        p.add_run("\nReferenced Sections:")

        for chunk in chunks:

            title = chunk.get("section_title", "")

            if title:
                p.add_run(
                    f"\n   • Chunk {chunk.get('chunk_number','')} ({title})"
                )
            else:
                p.add_run(
                    f"\n   • Chunk {chunk.get('chunk_number','')}"
                )

        logger.info(f"Creating hyperlink for {file_name}")
        logger.info(f"Source object: {first}")

        link = first.get("onedrive_url") or first.get("path")

        logger.info(f"Hyperlink target: {link}")

        logger.info(f"Creating hyperlink for {file_name}")
        logger.info(f"Target: {link}")

        if link:
            add_hyperlink(
                p,
                "Open Document",
                link
            )
        else:
            logger.warning(
                "No hyperlink available for %s",
                file_name
            )

        doc.add_paragraph()

    return doc